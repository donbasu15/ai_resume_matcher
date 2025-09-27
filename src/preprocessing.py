import re
import fitz  # PyMuPDF for PDF text extraction
from io import BytesIO
import spacy
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
import os

# Download NLTK data if not present
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Please install spaCy English model: python -m spacy download en_core_web_sm")
    nlp = None

def extract_text_from_pdf(file_input) -> str:
    """Extract text from PDF resume. Accepts file path string or file-like object."""
    if hasattr(file_input, 'read'):
        # It's a file-like object (Flask uploaded file)
        doc = fitz.open(stream=file_input.read(), filetype="pdf")
    else:
        # It's a file path string
        doc = fitz.open(file_input)
    
    text = " ".join([page.get_text() for page in doc])
    doc.close()
    return text

def extract_text_from_docx(file_input) -> str:
    """Extract text from DOCX resume. Accepts file path string or file-like object."""
    if hasattr(file_input, 'read'):
        # It's a file-like object (Flask uploaded file)
        doc = Document(file_input)
    else:
        # It's a file path string
        doc = Document(file_input)
    
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    return " ".join(text)

def extract_text_from_file(file_input, filename=None) -> str:
    """Extract text from various file formats."""
    if filename:
        ext = os.path.splitext(filename)[1].lower()
    elif hasattr(file_input, 'filename'):
        ext = os.path.splitext(file_input.filename)[1].lower()
    else:
        ext = '.pdf'  # default
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_input)
    elif ext == '.docx':
        return extract_text_from_docx(file_input)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def clean_text(text: str) -> str:
    """Basic cleaning of text."""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep letters, numbers, and basic punctuation
    text = re.sub(r'[^\w\s\-\.\,\(\)]', ' ', text)
    # Remove multiple spaces again
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def advanced_text_preprocessing(text: str, remove_stopwords: bool = False) -> str:
    """Advanced text preprocessing with lemmatization and optional stopword removal."""
    if not nlp:
        return clean_text(text).lower()
    
    # Clean text first
    text = clean_text(text)
    
    # Process with spaCy
    doc = nlp(text)
    
    # Extract lemmatized tokens
    tokens = []
    for token in doc:
        if not token.is_punct and not token.is_space:
            if remove_stopwords and token.is_stop:
                continue
            tokens.append(token.lemma_.lower())
    
    return " ".join(tokens)

def extract_sections(text: str) -> dict:
    """Extract different sections from resume text."""
    sections = {
        'contact': '',
        'summary': '',
        'experience': '',
        'education': '',
        'skills': '',
        'projects': '',
        'certifications': ''
    }
    
    # Define section patterns
    section_patterns = {
        'experience': r'(experience|work experience|employment|career|professional experience)',
        'education': r'(education|academic|degree|university|college|school)',
        'skills': r'(skills|technical skills|competencies|expertise|technologies)',
        'projects': r'(projects|personal projects|work projects)',
        'certifications': r'(certifications|certificates|licenses)',
        'summary': r'(summary|objective|profile|about)'
    }
    
    text_lower = text.lower()
    
    # Simple section extraction based on common headings
    for section, pattern in section_patterns.items():
        match = re.search(pattern, text_lower)
        if match:
            start_idx = match.start()
            # Find next section or end of text
            next_section_start = len(text)
            for other_pattern in section_patterns.values():
                if other_pattern != pattern:
                    next_match = re.search(other_pattern, text_lower[start_idx + 50:])
                    if next_match:
                        next_section_start = min(next_section_start, start_idx + 50 + next_match.start())
            
            sections[section] = text[start_idx:next_section_start].strip()
    
    return sections
